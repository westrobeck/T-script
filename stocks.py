import openai
import json
import os
import time
import re  # <-- Added this line
from docx import Document
import sys
import yfinance as yf
from datetime import datetime, timedelta
from scipy.stats import pearsonr  # For Correlation and p-value
# Replace with your API key (for testing only; in production, use environment variables)
openai.api_key = 

# Expanded list of company tickers (duplicates removed and as provided)
COMPANIES = [
    "BTSG", "DOCS", "EVH", "GDRX", "GKOS", "HQY", "ICUI", "INSP", "PEN", "PRCT", "RCM", "RGEN", "TEM", "TMDX",
    "NARI", "TNDM", "PRVA", "RXST", "TXG", "IRTC", "NVCR", "TDOC", "PGNY", "AHCO", "SDGR", "INMD", "PHR",
    "ESTA", "BLFS", "WEAV", "BFLY", "KIDS", "DH", "HCAT", "FNA", "QTRX", "TALK", "MXCT", "TCMD", "CERS",
    "VMD", "TMCI", "CLPT", "ACCD", "SGHT", "SMTI", "RCEL", "LUNG", "SOPH", "SMLR", "LFMD", "SENS", "NVRO",
    "BWAY", "CVRX", "AKYA", "ZOM", "TLSI", "OPRX", "MASS", "QIPT", "MYO", "CTSO", "TELA", "ECOR", "CCLD",
    "BNGO", "OM", "SSKN", "VERO", "RXRX", "SHC", "PACB", "LAB", "PROF", "CMPX", "ELAN", "SWTX", "MDAI",
    "SLDB", "IMNM", "ELTX"
]

# API configuration parameters updated for gpt-4o
MODEL = "gpt-4o"
TEMPERATURE = 0.83      
MAX_TOKENS = 1347   # Adjust as desired
TOP_P = 1.0
FREQUENCY_PENALTY = 0.0
PRESENCE_PENALTY = 0.0

# --- Earnings call dates dictionary ---
# Supply the actual earnings call date for each ticker (format: "YYYY-MM-DD")
earnings_dates = {
    "BTSG": "2024-11-01",
    "DOCS": "2025-02-06",
    "EVH": "2024-11-08",
    "GDRX": "2024-11-08",
    "GKOS": "2024-11-05",
    "HQY": "2024-12-10",
    "ICUI": "2024-11-13",
    "INSP": "2024-11-05",
    "PEN": "2024-10-31",
    "PRCT": "2024-10-29",
    "RCM": "2024-11-09",
    "RGEN": "2024-11-13",
    "TEM": "2024-11-05",
    "TMDX": "2024-11-29",
    "NARI": "2024-10-29",
    "TNDM": "2024-11-07",
    "PRVA": "2024-11-07",
    "RXST": "2024-11-07",
    "TXG": "2024-10-30",
    "IRTC": "2024-10-31",
    "NVCR": "2024-10-31",
    "TDOC": "2024-10-31",
    "PGNY": "2024-11-13",
    "AHCO": "2024-10-05",
    "SDGR": "2024-11-12",
    "INMD": "2024-10-30",
    "PHR": "2024-12-10",
    "ESTA": "2024-11-08",
    "BLFS": "2024-11-13",
    "WEAV": "2024-10-31",
    "BFLY": "2024-11-02",
    "KIDS": "2024-11-07",
    "DH": "2024-11-07",
    "HCAT": "2024-11-07",
    "FNA": "2024-11-13",
    "QTRX": "2024-11-13",
    "TALK": "2024-10-30",
    "MXCT": "2024-11-07",
    "TCMD": "2024-11-04",
    "CERS": "2024-10-31",
    "VMD": "2024-11-08",
    "TMCI": "2024-11-06",
    "CLPT": "2024-11-08",
    "ACCD": "2025-01-10",
    "SGHT": "2024-11-08",
    "SMTI": "2024-11-14",
    "RCEL": "2024-11-08",
    "LUNG": "2024-10-31",
    "SOPH": "2024-11-06",
    "SMLR": "2024-10-05",
    "LFMD": "2024-10-08",
    "SENS": "2024-11-08",
    "NVRO": "2024-11-12",
    "BWAY": "2024-11-13",
    "CVRX": "2024-10-30",
    "AKYA": "2024-11-15",
    "ZOM": "2024-11-08",
    "TLSI": "2024-11-15",
    "OPRX": "2024-11-14",
    "MASS": "2024-11-13",
    "QIPT": "2024-12-18",
    "MYO": "2024-11-07",
    "CTSO": "2024-11-08",
    "TELA": "2024-11-08",
    "ECOR": "2024-11-14",
    "CCLD": "2024-11-13",
    "BNGO": "2024-11-15",
    "OM": "2024-11-07",
    "SSKN": "2024-11-15",
    "VERO": "2024-11-14",
    "RXRX": "2024-11-13",
    "SHC": "2024-11-06",
    "PACB": "2024-11-09",
    "LAB": "2024-10-31",
    "PROF": "2024-11-09",
    "CMPX": "2024-12-05",
    "ELAN": "2024-11-08",
    "SWTX": "2024-11-13",
    "MDAI": "2024-11-07",
    "SLDB": "2024-11-07",
    "IMNM": "2024-11-14",
    "ELTX": "2024-11-14"
}

def clean_response(raw_text):
    """
    Remove markdown code block markers (e.g., triple backticks) if present.
    """
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        lines = cleaned.splitlines()
        if lines[0].strip().startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        cleaned = "\n".join(lines).strip()
    return cleaned

def analyze_transcript(file_path):
    """
    Reads a transcript file and uses the OpenAI API to perform a detailed, segmented sentiment analysis.
    The response must be returned in plain JSON format with exactly five keys:
      "score", "explanation", "investment_rating", "growth_adjectives", and "tone".
    """
    try:
        with open(file_path, "r", encoding="utf8") as f:
            transcript_text = f.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

    transcript_text = transcript_text.strip()
    file_size = os.path.getsize(file_path)
    print(f"DEBUG: {file_path} size: {file_size} bytes")
    if not transcript_text:
        print(f"WARNING: {file_path} appears to be empty after stripping.")
        return None
    else:
        print(f"DEBUG: {file_path} length: {len(transcript_text)} characters")
        print(f"DEBUG: First 300 characters: {transcript_text[:300]}")

    prompt = (
        "You are an expert equity research analyst known for your deep-dive analysis of earnings call transcripts. "
        "Analyze the following transcript thoroughly by performing these tasks:\n\n"
        "1. Break the transcript into key segments (e.g., revenue performance, margins, future guidance, and risk factors) "
        "and assess the positive and negative aspects of each segment.\n\n"
        "2. Based on this segmented evaluation, calculate an overall sentiment score on a scale from -1 (very negative) to 1 (very positive). "
        "If the transcript shows exceptional positive signals (e.g., extraordinary revenue growth and very strong guidance), "
        "assign a score above 0.8 and a 'Buy' rating. If it shows clear negative signals (e.g., significant revenue declines or high risks), "
        "assign a score below 0 and a 'Sell' rating. Otherwise, use a mid-range score and a 'Hold' rating.\n\n"
        "3. List at least two adjectives that describe the growth outlook.\n\n"
        "4. Identify the overall tone of the transcript (e.g., optimistic, cautious, aggressive, defensive).\n\n"
        "5. Provide a final investment recommendation of 'Buy', 'Hold', or 'Sell' with detailed supporting reasoning.\n\n"
        "Return your response in plain JSON format with exactly five keys: "
        "'score', 'explanation', 'investment_rating', 'growth_adjectives', and 'tone'. "
        "Do not include any markdown formatting or extra text.\n\n"
        "Transcript:\n" + transcript_text
    )

    max_retries = 5
    retries = 0
    while retries < max_retries:
        try:
            response = openai.ChatCompletion.create(
                model=MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=TEMPERATURE,
                max_tokens=MAX_TOKENS,
                top_p=TOP_P,
                frequency_penalty=FREQUENCY_PENALTY,
                presence_penalty=PRESENCE_PENALTY
            )
            break
        except Exception as e:
            error_str = str(e)
            if "Rate limit reached" in error_str:
                wait_time = 10
                match = re.search(r"Please try again in ([\d\.]+)s", error_str)
                if match:
                    wait_time = float(match.group(1)) + 1
                print(f"Rate limit error encountered for {file_path}. Waiting for {wait_time} seconds...")
                time.sleep(wait_time)
                retries += 1
            else:
                print(f"Error processing {file_path}: {e}")
                return None
    if retries == max_retries:
        print(f"Max retries reached for {file_path}. Skipping.")
        return None

    try:
        result_text = response.choices[0].message["content"].strip()
        print("DEBUG: Raw API response:")
        print(result_text)
        cleaned_result = clean_response(result_text)
        print("DEBUG: Cleaned API response:")
        print(cleaned_result)
        result = json.loads(cleaned_result)
        return result
    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return None

def backtest_portfolio(tickers, earnings_dates, period_days=30):
    """
    Given a list of tickers and an earnings date mapping, fetch historical closing prices using yfinance
    and compute the return from the earnings call date to earnings call date + period_days.
    Returns a dictionary of ticker returns.
    """
    returns = {}
    for ticker in tickers:
        if ticker not in earnings_dates:
            print(f"No earnings call date provided for {ticker}. Skipping backtest for this ticker.")
            continue
        start_date_str = earnings_dates[ticker]
        try:
            start_dt = datetime.strptime(start_date_str, "%Y-%m-%d").date()
        except Exception as e:
            print(f"Error parsing date for {ticker}: {e}")
            continue
        end_dt = start_dt + timedelta(days=period_days)
        start_date_str = start_dt.strftime("%Y-%m-%d")
        end_date_str = end_dt.strftime("%Y-%m-%d")
        try:
            data = yf.download(ticker, start=start_date_str, end=end_date_str, progress=False)
            if data.empty:
                print(f"No price data found for {ticker} between {start_date_str} and {end_date_str}.")
                continue
            start_price = float(data.iloc[0]['Close'])
            end_price = float(data.iloc[-1]['Close'])
            ret = (end_price - start_price) / start_price
            # If ret is a Series, extract its scalar value.
            if hasattr(ret, "item"):
                ret = ret.item()
            returns[ticker] = ret
        except Exception as e:
            print(f"Error fetching price data for {ticker}: {e}")
    return returns

def compute_correlation(results, earnings_dates, period_days=30):
    """
    For each ticker in results (if an earnings call date is provided), fetch the 1-month return,
    then compute the Pearson correlation between sentiment scores and monthly returns.
    Returns the correlation coefficient, p-value, and lists of sentiment scores and returns.
    """
    sentiment_scores = []
    monthly_returns = []
    for ticker, res in results.items():
        if ticker not in earnings_dates:
            continue
        start_date_str = earnings_dates[ticker]
        try:
            start_dt = datetime.strptime(start_date_str, "%Y-%m-%d").date()
        except Exception as e:
            print(f"Error parsing date for {ticker}: {e}")
            continue
        end_dt = start_dt + timedelta(days=period_days)
        start_date_str = start_dt.strftime("%Y-%m-%d")
        end_date_str = end_dt.strftime("%Y-%m-%d")
        try:
            data = yf.download(ticker, start=start_date_str, end=end_date_str, progress=False)
            if data.empty:
                print(f"No price data for {ticker} between {start_date_str} and {end_date_str}. Skipping.")
                continue
            start_price = float(data.iloc[0]['Close'])
            end_price = float(data.iloc[-1]['Close'])
            ret = (end_price - start_price) / start_price
            if hasattr(ret, "item"):
                ret = ret.item()
            sentiment_scores.append(res["score"])
            monthly_returns.append(ret)
        except Exception as e:
            print(f"Error fetching price data for {ticker}: {e}")
            continue

    if len(sentiment_scores) > 1 and len(monthly_returns) > 1:
        corr, p_value = pearsonr(sentiment_scores, monthly_returns)
        return corr, p_value, sentiment_scores, monthly_returns
    else:
        return None, None, sentiment_scores, monthly_returns

def main():
    transcripts_folder = "transcripts"
    results = {}

    print("Current working directory:", os.getcwd())

    # Loop over all .txt files in the transcripts folder
    for filename in os.listdir(transcripts_folder):
        if filename.endswith(".txt"):
            print(f"Found file: {filename}")
            ticker = filename.split("_")[0]
            print(f"Extracted ticker: {ticker}")
            if ticker in COMPANIES and ticker not in results:
                file_path = os.path.join(transcripts_folder, filename)
                print(f"Processing file with full path: {file_path}")
                analysis = analyze_transcript(file_path)
                if analysis:
                    explanation_value = analysis.get("explanation", "")
                    explanation_text = str(explanation_value).lower()
                    if "please provide the transcript" in explanation_text or "insufficient data" in explanation_text:
                        print(f"Transcript data missing for {filename}, skipping.")
                        continue
                    try:
                        score_val = analysis.get("score")
                        score = float(score_val) if score_val is not None else 0
                    except Exception as e:
                        score = 0
                    results[ticker] = {
                        "file": filename,
                        "ticker": ticker,
                        "score": score,
                        "explanation": analysis.get("explanation", "No explanation provided."),
                        "investment_rating": analysis.get("investment_rating", "No rating provided."),
                        "growth_adjectives": analysis.get("growth_adjectives", "No adjectives provided."),
                        "tone": analysis.get("tone", "No tone detected.")
                    }
                else:
                    print(f"Analysis for {filename} failed.")
                time.sleep(1)
            else:
                print(f"Ticker '{ticker}' is not in the list or already processed. Skipping {filename}.")

    if not results:
        print("No transcripts were successfully processed.")
        sys.exit(0)

    sorted_results = sorted(results.values(), key=lambda x: x["score"], reverse=True)

    print("\nRanked Transcripts by Sentiment Score:")
    for idx, res in enumerate(sorted_results, start=1):
        print(f"{idx}. File: {res['file']} | Ticker: {res['ticker']} | Score: {res['score']}")
        print("   Investment Rating:", res['investment_rating'])
        print("   Growth Adjectives:", res['growth_adjectives'])
        print("   Tone:", res['tone'])
        print("   Explanation:", res['explanation'])
        print("-" * 80)

    # Backtesting: select top 6 and bottom 6 companies for a 1-month period
    top_6 = sorted_results[:6]
    bottom_6 = sorted_results[-6:]
    top_tickers = [res['ticker'] for res in top_6]
    bottom_tickers = [res['ticker'] for res in bottom_6]

    print("\nBacktesting (1-Month) Results:")
    top_returns = backtest_portfolio(top_tickers, earnings_dates, period_days=30)
    bottom_returns = backtest_portfolio(bottom_tickers, earnings_dates, period_days=30)

    print("\nTop 6 (Long) Returns:")
    for ticker, ret in top_returns.items():
        # Force ret to be a float if it's not already
        r = ret.item() if hasattr(ret, "item") else float(ret)
        print(f"{ticker}: {r*100:.2f}%")
    print("Bottom 6 (Short) Returns:")
    for ticker, ret in bottom_returns.items():
        r = ret.item() if hasattr(ret, "item") else float(ret)
        print(f"{ticker}: {r*100:.2f}%")

    if top_returns and bottom_returns:
        avg_long_return = sum(top_returns.values()) / len(top_returns)
        avg_short_return = sum(bottom_returns.values()) / len(bottom_returns)
        # For a short position, a gain is when the price falls (i.e., negative return is good)
        portfolio_return = avg_long_return - avg_short_return
        print(f"\nSimple Portfolio Return (Long - Short): {portfolio_return*100:.2f}%")
    else:
        print("Not enough price data to compute portfolio return.")

    # Compute correlation between sentiment scores and 1-month returns for all tickers with valid data.
    corr, p_value, sentiments, monthly_returns = compute_correlation(results, earnings_dates, period_days=30)
    if corr is not None:
        print("\nCorrelation Analysis:")
        print(f"Pearson correlation coefficient between sentiment score and 1-month return: {corr:.4f}")
        print(f"P-value: {p_value:.4f}")
    else:
        print("Not enough data to compute correlation.")

    # Write transcript analysis results to a Word document
    document = Document()
    document.add_heading("Transcript Sentiment Analysis Results", 0)
    for idx, res in enumerate(sorted_results, start=1):
        document.add_heading(f"{idx}. {res['ticker']} - {res['file']}", level=1)
        document.add_paragraph(f"Sentiment Score: {res['score']}")
        document.add_paragraph(f"Investment Rating: {res['investment_rating']}")
        document.add_paragraph(f"Growth Adjectives: {res['growth_adjectives']}")
        document.add_paragraph(f"Tone: {res['tone']}")
        document.add_heading("Explanation", level=2)
        document.add_paragraph(str(res['explanation']))
        document.add_page_break()

    output_docx = "Transcript_Analysis_Results_All.docx"
    document.save(output_docx)
    print(f"\nResults have been written to {output_docx}")

if __name__ == "__main__":
    main()